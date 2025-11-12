"""
Generator für detaillierte Mathematik-Modul Implementierungsdokumentation
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    """Fügt formatierten Heading hinzu"""
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h

def add_paragraph(doc, text, bold=False, italic=False):
    """Fügt formatierten Paragraph hinzu"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    return p

def add_code_block(doc, code, language="python"):
    """Fügt Code-Block hinzu"""
    p = doc.add_paragraph(code)
    p.style = 'Normal'
    run = p.runs[0]
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    # Hellgrauer Hintergrund simulieren durch Einrückung
    p.paragraph_format.left_indent = Inches(0.5)
    return p

def add_bullet_list(doc, items):
    """Fügt Aufzählungsliste hinzu"""
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)

def add_numbered_list(doc, items):
    """Fügt nummerierte Liste hinzu"""
    for item in items:
        p = doc.add_paragraph(item, style='List Number')
        p.paragraph_format.left_indent = Inches(0.5)

def create_implementation_guide():
    """Erstellt die vollständige Implementierungsdokumentation"""
    doc = Document()

    # Titel
    title = doc.add_heading('KAI Mathematik-Modul', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph('Detaillierte Schritt-für-Schritt Implementierungsanleitung')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].italic = True

    doc.add_paragraph()
    add_paragraph(doc, 'Projekt: KAI (Konzeptueller AI Prototyp)', bold=True)
    add_paragraph(doc, 'Version: 1.0')
    add_paragraph(doc, 'Geschätzter Aufwand: 8-11 Tage')
    add_paragraph(doc, 'Neue Dateien: 7 | Modifizierte Dateien: 8')

    doc.add_page_break()

    # Inhaltsverzeichnis
    add_heading(doc, 'Inhaltsverzeichnis', 1)
    add_paragraph(doc, '1. Projektübersicht')
    add_paragraph(doc, '2. Architektur & Design')
    add_paragraph(doc, '3. Phase 1: Setup & Grundrechenarten (3-4 Tage)')
    add_paragraph(doc, '4. Phase 2: Konzepte & Eigenschaften (2-3 Tage)')
    add_paragraph(doc, '5. Phase 3: Erweiterte Mathematik (3-4 Tage)')
    add_paragraph(doc, '6. Testing-Strategie')
    add_paragraph(doc, '7. Validierung & Qualitätssicherung')
    add_paragraph(doc, '8. Checklisten')

    doc.add_page_break()

    # === ABSCHNITT 1: PROJEKTÜBERSICHT ===
    add_heading(doc, '1. Projektübersicht', 1)

    add_heading(doc, '1.1 Ziele', 2)
    add_bullet_list(doc, [
        'Arithmetisches Reasoning in KAI integrieren',
        'Natürlichsprachige Mathematik-Anfragen verarbeiten',
        'Bidirektionale Zahl-Wort-Konvertierung (deutsch)',
        'Transparente Proof-Trees für Berechnungen',
        'Erweiterbar für fortgeschrittene Mathematik'
    ])

    add_heading(doc, '1.2 Scope', 2)
    add_paragraph(doc, 'ENTHALTEN:', bold=True)
    add_bullet_list(doc, [
        'Grundrechenarten: +, -, *, /',
        'Vergleiche: <, >, =, ≤, ≥',
        'Zahlen-Eigenschaften: gerade/ungerade, Primzahl, Teiler',
        'Brüche und Dezimalzahlen',
        'Unbegrenzter Zahlenbereich (Python arbitrary precision)'
    ])

    add_paragraph(doc, 'NICHT ENTHALTEN (zukünftige Erweiterungen):', bold=True)
    add_bullet_list(doc, [
        'Algebraische Gleichungen',
        'Calculus (Ableitungen, Integrale)',
        'Lineare Algebra (Matrizen, Vektoren)',
        'Statistik & Wahrscheinlichkeit (außer Grundlagen)'
    ])

    add_heading(doc, '1.3 Erfolgskriterien', 2)
    add_numbered_list(doc, [
        'Alle Tests bestehen (100% Pass-Rate)',
        'Antwortzeit < 100ms für einfache Arithmetik',
        'Korrekte Proof-Trees für alle Operationen',
        'Natürlichsprachige Ein-/Ausgabe funktioniert',
        'Integration in bestehendes Reasoning-System nahtlos'
    ])

    doc.add_page_break()

    # === ABSCHNITT 2: ARCHITEKTUR & DESIGN ===
    add_heading(doc, '2. Architektur & Design', 1)

    add_heading(doc, '2.1 Systemarchitektur', 2)
    add_paragraph(doc, 'FLOW-DIAGRAMM:')
    add_code_block(doc, '''
User Input: "Was ist drei plus fünf?"
    ↓
[component_7_meaning_extractor]
    → Erkennt: ARITHMETIC_QUESTION (confidence=0.95)
    → Extrahiert: operation="addition", operands=["drei", "fünf"]
    ↓
[component_4_goal_planner]
    → Erstellt: Goal(PERFORM_CALCULATION, parameters={...})
    ↓
[kai_sub_goal_executor: ArithmeticStrategy]
    → Delegiert an ArithmeticEngine
    ↓
[component_52_arithmetic_reasoning: ArithmeticEngine]
    → 1. Konvertiert Wörter zu Zahlen (component_53)
    → 2. Führt Berechnung aus (3 + 5 = 8)
    → 3. Erstellt ProofTree
    → 4. Speichert in Neo4j (optional)
    ↓
[kai_response_formatter]
    → Konvertiert Zahl zu Wort: "acht"
    → Formatiert: "Die Summe von drei und fünf ist acht."
    ↓
User Output + Proof Tree (UI)
''')

    add_heading(doc, '2.2 Komponenten-Übersicht', 2)

    add_paragraph(doc, '1. component_52_arithmetic_reasoning.py (~1800 Zeilen)', bold=True)
    add_bullet_list(doc, [
        'ArithmeticEngine: Haupt-Orchestrator',
        'OperationRegistry: Registrierung aller Operationen',
        'BaseOperation: Abstract Base Class für Operationen',
        'Addition, Subtraction, Multiplication, Division: Konkrete Implementierungen',
        'ComparisonEngine: <, >, =, ≤, ≥',
        'PropertyChecker: gerade, ungerade, prim, etc.',
        'RationalArithmetic: Bruchrechnung',
        'ProofBuilder: Integration mit component_17'
    ])

    add_paragraph(doc, '2. component_53_number_language.py (~800 Zeilen)', bold=True)
    add_bullet_list(doc, [
        'NumberParser: Wort → Zahl (deutsch)',
        'NumberFormatter: Zahl → Wort (deutsch)',
        'NumberLearner: Neue Zahlen lernen und speichern',
        'Neo4j Integration: EQUIVALENT_TO Relationen',
        'Support: 0-999.999 (Basis), erweiterbar'
    ])

    add_heading(doc, '2.3 Datenmodell (Neo4j)', 2)
    add_paragraph(doc, 'NEUE NODE-TYPEN:', bold=True)
    add_code_block(doc, '''
(:Zahl {value: int/float, wort: str})
(:Operation {name: str, symbol: str, arity: int})
(:Eigenschaft {name: str, type: str})
''')

    add_paragraph(doc, 'NEUE RELATION-TYPEN:', bold=True)
    add_code_block(doc, '''
(:Zahl)-[:EQUIVALENT_TO]->(:Wort)           # 3 <-> "drei"
(:Zahl)-[:HAS_PROPERTY]->(:Eigenschaft)     # 4 -> "gerade"
(:Operation)-[:APPLIES_TO]->(:Zahl)         # Addition -> 3, 5
(:Zahl)-[:RESULT_OF {operation}]->(:Zahl)   # 8 RESULT_OF(+, 3, 5)
(:Zahl)-[:GREATER_THAN]->(:Zahl)            # Transitive Relation
''')

    doc.add_page_break()

    # === ABSCHNITT 3: PHASE 1 ===
    add_heading(doc, '3. Phase 1: Setup & Grundrechenarten (3-4 Tage)', 1)

    # Schritt 1
    add_heading(doc, 'Schritt 1.1: Modul-Skeleton erstellen', 2)
    add_paragraph(doc, 'DAUER: 4 Stunden', italic=True)

    add_paragraph(doc, 'ZIEL:', bold=True)
    add_paragraph(doc, 'Erstelle component_52_arithmetic_reasoning.py mit Basis-Struktur und Design Patterns.')

    add_paragraph(doc, 'IMPLEMENTIERUNG:', bold=True)
    add_code_block(doc, '''
# component_52_arithmetic_reasoning.py
"""
Arithmetisches Reasoning für KAI
Unterstützt: Grundrechenarten, Vergleiche, Eigenschaften, Brüche
"""

from abc import ABC, abstractmethod
from typing import Dict, List, Optional, Any, Tuple
from dataclasses import dataclass
from decimal import Decimal
from fractions import Fraction

from component_17_proof_explanation import ProofTree, ProofStep, StepType
from component_1_netzwerk_core import KonzeptNetzwerk


@dataclass
class ArithmeticResult:
    """Ergebnis einer arithmetischen Operation"""
    value: Any  # int, float, Fraction, Decimal
    proof_tree: ProofTree
    confidence: float = 1.0
    metadata: Dict[str, Any] = None


class BaseOperation(ABC):
    """Abstract Base Class für arithmetische Operationen"""

    def __init__(self, symbol: str, name: str, arity: int):
        self.symbol = symbol
        self.name = name
        self.arity = arity

    @abstractmethod
    def execute(self, *operands) -> ArithmeticResult:
        """Führt Operation aus und erstellt Proof"""
        pass

    @abstractmethod
    def validate(self, *operands) -> Tuple[bool, Optional[str]]:
        """Validiert Operanden (z.B. Division durch 0)"""
        pass


class OperationRegistry:
    """Registry für alle verfügbaren Operationen"""

    def __init__(self):
        self._operations: Dict[str, BaseOperation] = {}

    def register(self, operation: BaseOperation):
        """Registriert eine Operation"""
        self._operations[operation.symbol] = operation
        self._operations[operation.name] = operation

    def get(self, key: str) -> Optional[BaseOperation]:
        """Holt Operation nach Symbol oder Name"""
        return self._operations.get(key)

    def list_operations(self) -> List[str]:
        """Listet alle registrierten Operationen"""
        return list(set(self._operations.keys()))


class ArithmeticEngine:
    """Haupt-Engine für arithmetisches Reasoning"""

    def __init__(self, netzwerk: KonzeptNetzwerk):
        self.netzwerk = netzwerk
        self.registry = OperationRegistry()
        self._register_operations()

    def _register_operations(self):
        """Registriert alle Standard-Operationen"""
        self.registry.register(Addition())
        self.registry.register(Subtraction())
        self.registry.register(Multiplication())
        self.registry.register(Division())

    def calculate(self, operation: str, *operands) -> ArithmeticResult:
        """
        Führt Berechnung aus

        Args:
            operation: Operation (Symbol oder Name)
            operands: Operanden (bereits als Zahlen konvertiert)

        Returns:
            ArithmeticResult mit Wert, Proof und Confidence
        """
        op = self.registry.get(operation)
        if not op:
            raise ValueError(f"Unbekannte Operation: {operation}")

        # Validierung
        valid, error = op.validate(*operands)
        if not valid:
            raise ValueError(f"Validierung fehlgeschlagen: {error}")

        # Ausführung
        result = op.execute(*operands)

        # Optional: In Neo4j speichern
        self._persist_calculation(operation, operands, result)

        return result

    def _persist_calculation(self, operation: str, operands: tuple,
                            result: ArithmeticResult):
        """Speichert Berechnung in Neo4j (optional)"""
        # TODO: Implementieren
        pass
''')

    add_paragraph(doc, 'NÄCHSTE SCHRITTE:', bold=True)
    add_numbered_list(doc, [
        'Erstelle Datei component_52_arithmetic_reasoning.py',
        'Kopiere obigen Code',
        'Teste Import: python -c "from component_52_arithmetic_reasoning import ArithmeticEngine"',
        'Commit: "feat: Add ArithmeticEngine skeleton"'
    ])

    add_paragraph(doc, 'VALIDIERUNG:', bold=True)
    add_bullet_list(doc, [
        '✓ Datei existiert und ist importierbar',
        '✓ Alle Klassen haben Docstrings',
        '✓ Type Hints vorhanden',
        '✓ Code folgt PEP 8 (black/flake8)'
    ])

    doc.add_paragraph()

    # Schritt 2
    add_heading(doc, 'Schritt 1.2: Grundrechenarten implementieren', 2)
    add_paragraph(doc, 'DAUER: 6 Stunden', italic=True)

    add_paragraph(doc, 'ZIEL:', bold=True)
    add_paragraph(doc, 'Implementiere +, -, *, / mit Proof-Tree-Generation.')

    add_paragraph(doc, 'IMPLEMENTIERUNG:', bold=True)
    add_code_block(doc, '''
class Addition(BaseOperation):
    """Addition: a + b = c"""

    def __init__(self):
        super().__init__(symbol="+", name="addition", arity=2)

    def execute(self, a, b) -> ArithmeticResult:
        """Führt Addition aus"""
        result_value = a + b

        # Erstelle Proof Tree
        tree = ProofTree(conclusion=f"{a} + {b} = {result_value}")

        # Schritt 1: Operanden identifizieren
        step1 = ProofStep(
            step_type=StepType.GIVEN,
            description=f"Gegeben: Operanden {a} und {b}",
            confidence=1.0
        )
        tree.add_step(step1)

        # Schritt 2: Operation anwenden
        step2 = ProofStep(
            step_type=StepType.RULE_APPLICATION,
            description=f"Wende Addition an: {a} + {b}",
            rule="Arithmetik: Addition",
            confidence=1.0
        )
        tree.add_step(step2, parent=step1)

        # Schritt 3: Ergebnis
        step3 = ProofStep(
            step_type=StepType.CONCLUSION,
            description=f"Ergebnis: {result_value}",
            confidence=1.0
        )
        tree.add_step(step3, parent=step2)

        return ArithmeticResult(
            value=result_value,
            proof_tree=tree,
            confidence=1.0,
            metadata={"operation": "addition", "operands": [a, b]}
        )

    def validate(self, a, b) -> Tuple[bool, Optional[str]]:
        """Validiert Operanden"""
        if not isinstance(a, (int, float, Decimal, Fraction)):
            return False, f"Operand a ist keine Zahl: {type(a)}"
        if not isinstance(b, (int, float, Decimal, Fraction)):
            return False, f"Operand b ist keine Zahl: {type(b)}"
        return True, None


class Division(BaseOperation):
    """Division: a / b = c (mit Division-durch-0-Check)"""

    def __init__(self):
        super().__init__(symbol="/", name="division", arity=2)

    def execute(self, a, b) -> ArithmeticResult:
        """Führt Division aus"""
        # Verwende Fraction für exakte Brüche
        result_value = Fraction(a, b)

        tree = ProofTree(conclusion=f"{a} / {b} = {result_value}")

        step1 = ProofStep(
            step_type=StepType.GIVEN,
            description=f"Gegeben: Dividend {a}, Divisor {b}",
            confidence=1.0
        )
        tree.add_step(step1)

        step2 = ProofStep(
            step_type=StepType.CONSTRAINT_CHECK,
            description=f"Prüfe: {b} ≠ 0",
            confidence=1.0
        )
        tree.add_step(step2, parent=step1)

        step3 = ProofStep(
            step_type=StepType.RULE_APPLICATION,
            description=f"Wende Division an: {a} / {b}",
            rule="Arithmetik: Division",
            confidence=1.0
        )
        tree.add_step(step3, parent=step2)

        step4 = ProofStep(
            step_type=StepType.CONCLUSION,
            description=f"Ergebnis: {result_value} (exakter Bruch)",
            confidence=1.0
        )
        tree.add_step(step4, parent=step3)

        return ArithmeticResult(
            value=result_value,
            proof_tree=tree,
            confidence=1.0,
            metadata={"operation": "division", "operands": [a, b]}
        )

    def validate(self, a, b) -> Tuple[bool, Optional[str]]:
        """Validiert Operanden (Division durch 0!)"""
        if not isinstance(a, (int, float, Decimal, Fraction)):
            return False, f"Dividend ist keine Zahl: {type(a)}"
        if not isinstance(b, (int, float, Decimal, Fraction)):
            return False, f"Divisor ist keine Zahl: {type(b)}"
        if b == 0:
            return False, "Division durch 0 ist nicht definiert"
        return True, None
''')

    add_paragraph(doc, 'AUFGABEN:', bold=True)
    add_numbered_list(doc, [
        'Implementiere Addition (siehe oben)',
        'Implementiere Subtraction (analog zu Addition)',
        'Implementiere Multiplication (analog)',
        'Implementiere Division (siehe oben, wichtig: Division-durch-0!)',
        'Registriere alle in ArithmeticEngine._register_operations()',
        'Teste manuell: engine.calculate("+", 3, 5)'
    ])

    add_paragraph(doc, 'VALIDIERUNG:', bold=True)
    add_bullet_list(doc, [
        '✓ Alle 4 Operationen implementiert',
        '✓ Proof Trees korrekt generiert',
        '✓ Division durch 0 wirft ValueError',
        '✓ Fraction für exakte Division verwendet'
    ])

    doc.add_paragraph()

    # Schritt 3
    add_heading(doc, 'Schritt 1.3: Zahl-Wort-System (component_53)', 2)
    add_paragraph(doc, 'DAUER: 8 Stunden', italic=True)

    add_paragraph(doc, 'ZIEL:', bold=True)
    add_paragraph(doc, 'Bidirektionale Konvertierung zwischen Zahlen und deutschen Wörtern.')

    add_paragraph(doc, 'IMPLEMENTIERUNG (Auszug):', bold=True)
    add_code_block(doc, '''
# component_53_number_language.py
"""
Zahl-Wort-Konvertierung für KAI (Deutsch)
Unterstützt: 0-999.999, erweiterbar auf beliebig große Zahlen
"""

from typing import Optional, Dict
from component_1_netzwerk_core import KonzeptNetzwerk


class NumberParser:
    """Konvertiert deutsche Zahlwörter zu Zahlen"""

    # Basis-Mapping (1-20)
    BASIC_NUMBERS = {
        "null": 0, "eins": 1, "zwei": 2, "drei": 3, "vier": 4,
        "fünf": 5, "sechs": 6, "sieben": 7, "acht": 8, "neun": 9,
        "zehn": 10, "elf": 11, "zwölf": 12, "dreizehn": 13,
        "vierzehn": 14, "fünfzehn": 15, "sechzehn": 16,
        "siebzehn": 17, "achtzehn": 18, "neunzehn": 19, "zwanzig": 20
    }

    # Zehner
    TENS = {
        "zwanzig": 20, "dreißig": 30, "vierzig": 40, "fünfzig": 50,
        "sechzig": 60, "siebzig": 70, "achtzig": 80, "neunzig": 90
    }

    # Größere Einheiten
    MAGNITUDES = {
        "hundert": 100,
        "tausend": 1000,
        "million": 1000000,
        "milliarde": 1000000000
    }

    def __init__(self, netzwerk: Optional[KonzeptNetzwerk] = None):
        self.netzwerk = netzwerk
        self._load_learned_numbers()

    def _load_learned_numbers(self):
        """Lädt gelernte Zahlen aus Neo4j"""
        if not self.netzwerk:
            return

        # TODO: Query für EQUIVALENT_TO Relationen
        pass

    def parse(self, word: str) -> Optional[int]:
        """
        Konvertiert Wort zu Zahl

        Beispiele:
            "drei" → 3
            "einundzwanzig" → 21
            "zweihundertfünfundvierzig" → 245
            "dreitausendzweihundert" → 3200
        """
        word = word.lower().strip()

        # Direkt-Lookup
        if word in self.BASIC_NUMBERS:
            return self.BASIC_NUMBERS[word]

        # Komplexe Zahlen parsen
        return self._parse_complex(word)

    def _parse_complex(self, word: str) -> Optional[int]:
        """Parst zusammengesetzte Zahlen"""
        total = 0
        current = 0

        # Split bei "und" (einundzwanzig → ein + zwanzig)
        if "und" in word:
            parts = word.split("und")
            if len(parts) == 2:
                ones = self.BASIC_NUMBERS.get(parts[0])
                tens = self.TENS.get(parts[1])
                if ones and tens:
                    return tens + ones

        # Weitere Parsing-Logik für hundert, tausend, etc.
        # TODO: Vollständige Implementierung

        return None


class NumberFormatter:
    """Konvertiert Zahlen zu deutschen Wörtern"""

    def format(self, number: int) -> str:
        """
        Konvertiert Zahl zu Wort

        Beispiele:
            3 → "drei"
            21 → "einundzwanzig"
            245 → "zweihundertfünfundvierzig"
        """
        if number == 0:
            return "null"

        if number < 0:
            return "minus " + self.format(abs(number))

        if number <= 20:
            return self._get_basic_word(number)

        if number < 100:
            return self._format_tens(number)

        if number < 1000:
            return self._format_hundreds(number)

        # Weitere Formatierung für größere Zahlen
        return str(number)  # Fallback

    def _format_tens(self, number: int) -> str:
        """Formatiert Zehner (21-99)"""
        tens = (number // 10) * 10
        ones = number % 10

        tens_word = self._get_tens_word(tens)

        if ones == 0:
            return tens_word
        else:
            ones_word = self._get_basic_word(ones)
            return f"{ones_word}und{tens_word}"

    # ... weitere Methoden
''')

    add_paragraph(doc, 'AUFGABEN:', bold=True)
    add_numbered_list(doc, [
        'Erstelle component_53_number_language.py',
        'Implementiere NumberParser.parse() vollständig (0-999)',
        'Implementiere NumberFormatter.format() vollständig (0-999)',
        'Teste Roundtrip: parse(format(n)) == n für alle n in 0-999',
        'Implementiere Neo4j Integration (EQUIVALENT_TO Relationen)',
        'Füge Basis-Zahlen zu setup_initial_knowledge.py hinzu'
    ])

    add_paragraph(doc, 'VALIDIERUNG:', bold=True)
    add_bullet_list(doc, [
        '✓ Alle Zahlen 0-20 korrekt konvertiert',
        '✓ Zehner (21-99) korrekt',
        '✓ Hunderter (100-999) korrekt',
        '✓ Roundtrip-Test bestanden',
        '✓ Neo4j Integration funktioniert'
    ])

    doc.add_page_break()

    # Schritt 4: Intent Detection
    add_heading(doc, 'Schritt 1.4: Intent Detection erweitern', 2)
    add_paragraph(doc, 'DAUER: 4 Stunden', italic=True)

    add_paragraph(doc, 'ZIEL:', bold=True)
    add_paragraph(doc, 'Erkenne arithmetische Fragen in component_7_meaning_extractor.py')

    add_paragraph(doc, 'ÄNDERUNGEN:', bold=True)
    add_paragraph(doc, '1. Füge neues MeaningPoint zu component_5_linguistik_strukturen.py:')
    add_code_block(doc, '''
class MeaningPoint(Enum):
    QUESTION = "frage"
    COMMAND = "befehl"
    DEFINITION = "definition"
    TASK = "aufgabe"
    ARITHMETIC_QUESTION = "arithmetic_frage"  # NEU
''')

    add_paragraph(doc, '2. Erweitere component_7_meaning_extractor.py:')
    add_code_block(doc, '''
class MeaningExtractor:
    # ...

    def _detect_arithmetic_question(self, text: str, doc) -> float:
        """
        Erkennt arithmetische Fragen

        Patterns:
            "Was ist 3 + 5?"          → 0.95
            "Wie viel ist 7 mal 8?"   → 0.93
            "Wieviel sind 10 durch 2?" → 0.92
            "Berechne 15 minus 6"     → 0.90
        """
        text_lower = text.lower()

        # Arithmetische Trigger-Wörter
        arithmetic_triggers = [
            "plus", "minus", "mal", "geteilt", "durch",
            "+", "-", "*", "/", "×", "÷"
        ]

        # Frage-Trigger
        question_triggers = ["was ist", "wie viel", "wieviel", "berechne"]

        has_arithmetic = any(trigger in text_lower for trigger in arithmetic_triggers)
        has_question = any(trigger in text_lower for trigger in question_triggers)

        if has_arithmetic and has_question:
            return 0.95
        elif has_arithmetic and text.endswith("?"):
            return 0.90
        elif has_arithmetic:
            return 0.80
        else:
            return 0.0

    def extract_meaning(self, text: str) -> List[ExtractedMeaning]:
        # Bestehender Code ...

        # NEU: Prüfe auf arithmetische Frage
        arith_score = self._detect_arithmetic_question(text, doc)
        if arith_score > 0.7:
            meanings.append(ExtractedMeaning(
                meaning=MeaningPoint.ARITHMETIC_QUESTION,
                confidence=arith_score,
                subjects=[],  # Wird später von Parser extrahiert
                predicates=[],
                objects=[]
            ))
''')

    add_paragraph(doc, 'AUFGABEN:', bold=True)
    add_numbered_list(doc, [
        'Füge ARITHMETIC_QUESTION zu MeaningPoint hinzu',
        'Implementiere _detect_arithmetic_question() in component_7',
        'Integriere in extract_meaning() Pipeline',
        'Teste mit: "Was ist drei plus fünf?"',
        'Validiere, dass bestehende Tests noch laufen'
    ])

    doc.add_paragraph()

    # Schritt 5: Integration
    add_heading(doc, 'Schritt 1.5: Integration in Reasoning-System', 2)
    add_paragraph(doc, 'DAUER: 6 Stunden', italic=True)

    add_paragraph(doc, 'ZIEL:', bold=True)
    add_paragraph(doc, 'Integriere Arithmetik in Goal Planner, Sub-Goal Executor und Reasoning Orchestrator.')

    add_paragraph(doc, 'ÄNDERUNGEN:', bold=True)

    add_paragraph(doc, '1. component_4_goal_planner.py - Neuer Goal-Type:')
    add_code_block(doc, '''
class GoalType(Enum):
    ANSWER_QUESTION = "answer_question"
    LEARN_KNOWLEDGE = "learn_knowledge"
    PERFORM_TASK = "perform_task"
    PERFORM_CALCULATION = "perform_calculation"  # NEU
    # ...

class GoalPlanner:
    def _plan_for_arithmetic(self, meaning: ExtractedMeaning) -> Goal:
        """Erstellt Goal für arithmetische Berechnung"""
        return Goal(
            goal_type=GoalType.PERFORM_CALCULATION,
            parameters={
                "text": meaning.original_text,
                "confidence": meaning.confidence
            },
            priority=0.9
        )
''')

    add_paragraph(doc, '2. kai_sub_goal_executor.py - Neue Strategy:')
    add_code_block(doc, '''
class ArithmeticStrategy(Strategy):
    """Strategie für arithmetische Berechnungen"""

    def __init__(self, netzwerk: KonzeptNetzwerk):
        self.netzwerk = netzwerk
        self.engine = ArithmeticEngine(netzwerk)
        self.parser = NumberParser(netzwerk)
        self.formatter = NumberFormatter(netzwerk)

    def execute(self, goal: Goal, context: Dict) -> StrategyResult:
        text = goal.parameters["text"]

        # 1. Parse Text → Operation + Operanden
        operation, operands_words = self._parse_arithmetic_text(text)

        # 2. Konvertiere Wörter zu Zahlen
        operands = [self.parser.parse(word) for word in operands_words]

        # 3. Führe Berechnung aus
        result = self.engine.calculate(operation, *operands)

        # 4. Formatiere Ergebnis
        result_word = self.formatter.format(result.value)

        return StrategyResult(
            success=True,
            result=result_word,
            confidence=result.confidence,
            proof_tree=result.proof_tree
        )

# In SubGoalExecutor.__init__():
self.strategies[GoalType.PERFORM_CALCULATION] = ArithmeticStrategy(netzwerk)
''')

    add_paragraph(doc, '3. kai_reasoning_orchestrator.py - Register Arithmetic:')
    add_code_block(doc, '''
class ReasoningOrchestrator:
    def __init__(self, netzwerk):
        # ...
        self.arithmetic_engine = ArithmeticEngine(netzwerk)  # NEU

    def reason(self, question: str, context: Dict) -> ReasoningResult:
        # Prüfe ob arithmetische Frage
        if self._is_arithmetic(question):
            return self._arithmetic_strategy(question, context)

        # Bestehende Strategien...

    def _arithmetic_strategy(self, question: str, context: Dict) -> ReasoningResult:
        """8. Strategie: Arithmetisches Reasoning"""
        # Delegiere an ArithmeticStrategy
        # ...
''')

    add_paragraph(doc, 'AUFGABEN:', bold=True)
    add_numbered_list(doc, [
        'Füge PERFORM_CALCULATION zu GoalType hinzu',
        'Implementiere ArithmeticStrategy in kai_sub_goal_executor.py',
        'Registriere Strategy in SubGoalExecutor.__init__()',
        'Erweitere ReasoningOrchestrator um Arithmetik',
        'End-to-End Test: "Was ist drei plus fünf?" → "acht"'
    ])

    doc.add_paragraph()

    # Schritt 6: Tests
    add_heading(doc, 'Schritt 1.6: Tests für Phase 1', 2)
    add_paragraph(doc, 'DAUER: 6 Stunden', italic=True)

    add_paragraph(doc, 'ZIEL:', bold=True)
    add_paragraph(doc, 'Erstelle umfassende Test-Suite für Grundrechenarten.')

    add_paragraph(doc, 'TEST-STRUKTUR:', bold=True)
    add_code_block(doc, '''
# tests/test_arithmetic_basic.py
import pytest
from component_52_arithmetic_reasoning import ArithmeticEngine, Addition, Division
from component_53_number_language import NumberParser, NumberFormatter
from component_1_netzwerk_core import KonzeptNetzwerk

class TestArithmeticEngine:
    """Tests für ArithmeticEngine"""

    @pytest.fixture
    def engine(self):
        netzwerk = KonzeptNetzwerk(uri="bolt://localhost:7687",
                                   user="neo4j", password="password")
        return ArithmeticEngine(netzwerk)

    def test_addition_simple(self, engine):
        result = engine.calculate("+", 3, 5)
        assert result.value == 8
        assert result.confidence == 1.0
        assert result.proof_tree is not None

    def test_division_by_zero(self, engine):
        with pytest.raises(ValueError, match="Division durch 0"):
            engine.calculate("/", 10, 0)

    def test_proof_tree_structure(self, engine):
        result = engine.calculate("*", 7, 8)
        tree = result.proof_tree
        assert tree.conclusion == "7 * 8 = 56"
        assert len(tree.steps) >= 3  # GIVEN, RULE_APPLICATION, CONCLUSION

class TestNumberParser:
    """Tests für Zahl-Wort-Konvertierung"""

    @pytest.fixture
    def parser(self):
        return NumberParser()

    @pytest.mark.parametrize("word,expected", [
        ("null", 0),
        ("eins", 1),
        ("zehn", 10),
        ("zwanzig", 20),
        ("einundzwanzig", 21),
        ("neunundneunzig", 99),
        ("hundert", 100),
        ("zweihundertfünfundvierzig", 245),
    ])
    def test_parse_numbers(self, parser, word, expected):
        assert parser.parse(word) == expected

    def test_roundtrip(self, parser):
        formatter = NumberFormatter()
        for n in range(0, 100):
            word = formatter.format(n)
            parsed = parser.parse(word)
            assert parsed == n, f"Roundtrip failed for {n}: {word} -> {parsed}"

# Weitere Test-Klassen...
''')

    add_paragraph(doc, 'TEST-KATEGORIEN:', bold=True)
    add_numbered_list(doc, [
        'TestArithmeticEngine: Alle 4 Grundrechenarten, Edge Cases',
        'TestNumberParser: 0-999, Roundtrip, Fehlerbehandlung',
        'TestIntegration: End-to-End mit UI-Input',
        'TestProofTrees: Struktur, Confidence, Metadaten',
        'TestNeo4jIntegration: Persistierung, Abfragen'
    ])

    add_paragraph(doc, 'ZIEL: 20+ Tests, 100% Pass-Rate', bold=True)

    doc.add_page_break()

    # === ABSCHNITT 4: PHASE 2 ===
    add_heading(doc, '4. Phase 2: Konzepte & Eigenschaften (2-3 Tage)', 1)

    add_heading(doc, 'Schritt 2.1: Vergleiche implementieren', 2)
    add_paragraph(doc, 'DAUER: 4 Stunden', italic=True)

    add_paragraph(doc, 'ZIEL:', bold=True)
    add_paragraph(doc, 'Implementiere <, >, =, ≤, ≥ mit transitiven Schlüssen.')

    add_paragraph(doc, 'IMPLEMENTIERUNG:', bold=True)
    add_code_block(doc, '''
# In component_52_arithmetic_reasoning.py

class ComparisonEngine:
    """Engine für Vergleichsoperationen"""

    def __init__(self, netzwerk: KonzeptNetzwerk):
        self.netzwerk = netzwerk

    def compare(self, a, b, operator: str) -> ArithmeticResult:
        """
        Vergleicht zwei Zahlen

        Args:
            a, b: Zahlen
            operator: "<", ">", "=", "<=", ">="

        Returns:
            ArithmeticResult mit bool-Wert
        """
        operations = {
            "<": lambda x, y: x < y,
            ">": lambda x, y: x > y,
            "=": lambda x, y: x == y,
            "<=": lambda x, y: x <= y,
            ">=": lambda x, y: x >= y
        }

        if operator not in operations:
            raise ValueError(f"Unbekannter Vergleichsoperator: {operator}")

        result_value = operations[operator](a, b)

        # Proof Tree
        tree = ProofTree(conclusion=f"{a} {operator} {b} ist {result_value}")

        step1 = ProofStep(
            step_type=StepType.GIVEN,
            description=f"Gegeben: {a} und {b}",
            confidence=1.0
        )
        tree.add_step(step1)

        step2 = ProofStep(
            step_type=StepType.RULE_APPLICATION,
            description=f"Vergleiche {a} {operator} {b}",
            rule=f"Arithmetik: Vergleich {operator}",
            confidence=1.0
        )
        tree.add_step(step2, parent=step1)

        step3 = ProofStep(
            step_type=StepType.CONCLUSION,
            description=f"Ergebnis: {result_value}",
            confidence=1.0
        )
        tree.add_step(step3, parent=step2)

        return ArithmeticResult(
            value=result_value,
            proof_tree=tree,
            confidence=1.0,
            metadata={"operation": "comparison", "operator": operator}
        )

    def transitive_inference(self, relations: List[Tuple]) -> List[Tuple]:
        """
        Leitet transitive Relationen ab

        Beispiel:
            Input: [(3, "<", 5), (5, "<", 7)]
            Output: [(3, "<", 7)]
        """
        # Implementiere transitive Schließung
        # A < B ∧ B < C → A < C
        inferred = []

        for i, (a1, op1, b1) in enumerate(relations):
            for j, (a2, op2, b2) in enumerate(relations):
                if i != j and b1 == a2 and op1 == op2 == "<":
                    inferred.append((a1, "<", b2))

        return inferred
''')

    add_paragraph(doc, 'AUFGABEN:', bold=True)
    add_numbered_list(doc, [
        'Implementiere ComparisonEngine',
        'Implementiere alle 5 Vergleichsoperatoren',
        'Implementiere transitive_inference()',
        'Integriere in ArithmeticEngine',
        'Teste Transitivität: 3<5, 5<7 → 3<7'
    ])

    doc.add_paragraph()

    add_heading(doc, 'Schritt 2.2: Zahlen-Eigenschaften', 2)
    add_paragraph(doc, 'DAUER: 6 Stunden', italic=True)

    add_paragraph(doc, 'ZIEL:', bold=True)
    add_paragraph(doc, 'Implementiere Eigenschafts-Checks: gerade, ungerade, prim, Teiler.')

    add_paragraph(doc, 'IMPLEMENTIERUNG:', bold=True)
    add_code_block(doc, '''
class PropertyChecker:
    """Prüft Eigenschaften von Zahlen"""

    def __init__(self, netzwerk: KonzeptNetzwerk):
        self.netzwerk = netzwerk

    def is_even(self, n: int) -> ArithmeticResult:
        """Prüft ob Zahl gerade ist"""
        result_value = (n % 2 == 0)

        tree = ProofTree(conclusion=f"{n} ist {'gerade' if result_value else 'ungerade'}")

        step1 = ProofStep(
            step_type=StepType.GIVEN,
            description=f"Gegeben: {n}",
            confidence=1.0
        )
        tree.add_step(step1)

        step2 = ProofStep(
            step_type=StepType.RULE_APPLICATION,
            description=f"Prüfe: {n} % 2 = {n % 2}",
            rule="Definition: n ist gerade ⟺ n % 2 = 0",
            confidence=1.0
        )
        tree.add_step(step2, parent=step1)

        step3 = ProofStep(
            step_type=StepType.CONCLUSION,
            description=f"{n} ist {'gerade' if result_value else 'ungerade'}",
            confidence=1.0
        )
        tree.add_step(step3, parent=step2)

        return ArithmeticResult(value=result_value, proof_tree=tree, confidence=1.0)

    def is_prime(self, n: int) -> ArithmeticResult:
        """Prüft ob Zahl prim ist"""
        if n < 2:
            return ArithmeticResult(value=False, proof_tree=self._build_prime_proof(n, False))

        for i in range(2, int(n**0.5) + 1):
            if n % i == 0:
                return ArithmeticResult(value=False, proof_tree=self._build_prime_proof(n, False, divisor=i))

        return ArithmeticResult(value=True, proof_tree=self._build_prime_proof(n, True))

    def find_divisors(self, n: int) -> ArithmeticResult:
        """Findet alle Teiler einer Zahl"""
        divisors = [i for i in range(1, n + 1) if n % i == 0]

        tree = ProofTree(conclusion=f"Teiler von {n}: {divisors}")
        # ... Proof Steps

        return ArithmeticResult(value=divisors, proof_tree=tree, confidence=1.0)

    def _build_prime_proof(self, n: int, is_prime: bool, divisor: int = None) -> ProofTree:
        """Erstellt Proof Tree für Primzahl-Check"""
        tree = ProofTree(conclusion=f"{n} ist {'prim' if is_prime else 'nicht prim'}")
        # ... detaillierte Schritte
        return tree
''')

    add_paragraph(doc, 'AUFGABEN:', bold=True)
    add_numbered_list(doc, [
        'Implementiere PropertyChecker mit is_even(), is_odd()',
        'Implementiere is_prime() mit effizienter Methode',
        'Implementiere find_divisors()',
        'Speichere Eigenschaften in Neo4j (HAS_PROPERTY)',
        'Teste mit Primzahlen: 2, 3, 5, 7, 11, ... (nicht: 4, 6, 8, 9, ...)'
    ])

    doc.add_paragraph()

    add_heading(doc, 'Schritt 2.3: Konzepte als abstrakte Operationen', 2)
    add_paragraph(doc, 'DAUER: 4 Stunden', italic=True)

    add_paragraph(doc, 'ZIEL:', bold=True)
    add_paragraph(doc, 'Modelliere "Summe", "Produkt", "Differenz" als eigenständige Konzepte.')

    add_paragraph(doc, 'BEISPIEL:', bold=True)
    add_paragraph(doc, 'Frage: "Was ist die Summe von drei und fünf?"')
    add_paragraph(doc, 'KAI versteht: "Summe" = Konzept, das Addition repräsentiert')
    add_paragraph(doc, 'Antwort: "Die Summe ist acht."')

    add_paragraph(doc, 'IMPLEMENTIERUNG:', bold=True)
    add_code_block(doc, '''
# In component_7_meaning_extractor.py erweitern

ARITHMETIC_CONCEPTS = {
    "summe": ("addition", "+"),
    "differenz": ("subtraction", "-"),
    "produkt": ("multiplication", "*"),
    "quotient": ("division", "/")
}

def _extract_arithmetic_concept(self, text: str) -> Optional[Tuple[str, str]]:
    """Extrahiert arithmetisches Konzept aus Text"""
    text_lower = text.lower()

    for concept, (operation, symbol) in ARITHMETIC_CONCEPTS.items():
        if concept in text_lower:
            return operation, symbol

    return None
''')

    add_paragraph(doc, 'Neo4j Modellierung:')
    add_code_block(doc, '''
(:Konzept {name: "Summe"})-[:EQUIVALENT_TO]->(:Operation {name: "Addition"})
(:Konzept {name: "Produkt"})-[:EQUIVALENT_TO]->(:Operation {name: "Multiplikation"})
''')

    doc.add_page_break()

    # === ABSCHNITT 5: PHASE 3 ===
    add_heading(doc, '5. Phase 3: Erweiterte Mathematik (3-4 Tage)', 1)

    add_heading(doc, 'Schritt 3.1: Brüche (Rational Arithmetic)', 2)
    add_paragraph(doc, 'DAUER: 6 Stunden', italic=True)

    add_paragraph(doc, 'ZIEL:', bold=True)
    add_paragraph(doc, 'Unterstütze Bruchrechnung mit exakten Ergebnissen (Python fractions).')

    add_paragraph(doc, 'IMPLEMENTIERUNG:', bold=True)
    add_code_block(doc, '''
from fractions import Fraction
import math

class RationalArithmetic:
    """Bruchrechnung"""

    def __init__(self):
        pass

    def add(self, a: Fraction, b: Fraction) -> Fraction:
        """Addiert zwei Brüche"""
        return a + b  # Python Fraction unterstützt + direkt

    def multiply(self, a: Fraction, b: Fraction) -> Fraction:
        """Multipliziert zwei Brüche"""
        return a * b

    def simplify(self, fraction: Fraction) -> Fraction:
        """Kürzt Bruch (automatisch durch Fraction)"""
        return fraction

    def to_mixed_number(self, fraction: Fraction) -> Tuple[int, Fraction]:
        """
        Konvertiert zu gemischter Zahl

        Beispiel: 7/3 → 2 und 1/3
        """
        whole = fraction.numerator // fraction.denominator
        remainder = fraction.numerator % fraction.denominator
        return whole, Fraction(remainder, fraction.denominator)

    def gcd(self, a: int, b: int) -> int:
        """Größter gemeinsamer Teiler"""
        return math.gcd(a, b)

    def lcm(self, a: int, b: int) -> int:
        """Kleinstes gemeinsames Vielfaches"""
        return abs(a * b) // math.gcd(a, b)
''')

    add_paragraph(doc, 'Integration:')
    add_code_block(doc, '''
# In ArithmeticEngine
def calculate(self, operation: str, *operands):
    # Konvertiere operands zu Fraction falls nötig
    operands_rational = [Fraction(op) if isinstance(op, (int, float)) else op
                        for op in operands]

    # Führe Operation aus
    result = op.execute(*operands_rational)

    # Formatiere Ergebnis
    if isinstance(result.value, Fraction):
        if result.value.denominator == 1:
            result.value = result.value.numerator  # Ganzzahl
        else:
            # Behalte als Bruch
            pass

    return result
''')

    doc.add_paragraph()

    add_heading(doc, 'Schritt 3.2: Dezimalzahlen & Rundung', 2)
    add_paragraph(doc, 'DAUER: 4 Stunden', italic=True)

    add_paragraph(doc, 'ZIEL:', bold=True)
    add_paragraph(doc, 'Unterstütze Float-Arithmetik mit konfigurierbarer Präzision.')

    add_paragraph(doc, 'IMPLEMENTIERUNG:', bold=True)
    add_code_block(doc, '''
from decimal import Decimal, getcontext

class DecimalArithmetic:
    """Dezimalzahl-Arithmetik mit konfigurierbarer Präzision"""

    def __init__(self, precision: int = 10):
        self.precision = precision
        getcontext().prec = precision

    def calculate(self, operation: str, *operands):
        """Führt Operation mit Decimal-Präzision aus"""
        operands_decimal = [Decimal(str(op)) for op in operands]
        # ... Operation ausführen

    def round(self, value: float, decimals: int) -> float:
        """Rundet auf n Dezimalstellen"""
        return round(value, decimals)
''')

    doc.add_paragraph()

    add_heading(doc, 'Schritt 3.3: UI-Integration (Math Proof Tree)', 2)
    add_paragraph(doc, 'DAUER: 6 Stunden', italic=True)

    add_paragraph(doc, 'ZIEL:', bold=True)
    add_paragraph(doc, 'Zeige mathematische Proof Trees in der GUI an.')

    add_paragraph(doc, 'ÄNDERUNGEN:', bold=True)
    add_paragraph(doc, '1. main_ui_graphical.py erweitern:')
    add_code_block(doc, '''
class MainWindow:
    def __init__(self):
        # ...

        # Neuer Tab für Math Proof Tree (optional, oder in bestehendem Tab)
        self.math_proof_tree_widget = ProofTreeWidget()
        self.tabs.addTab(self.math_proof_tree_widget, "Math Reasoning")

    def update_proof_tree(self, proof_tree: ProofTree):
        """Aktualisiert Proof Tree Widget"""
        self.math_proof_tree_widget.display_proof(proof_tree)
''')

    add_paragraph(doc, '2. Formatierung für mathematische Proofs:')
    add_code_block(doc, '''
# Proof Tree Rendering mit mathematischen Symbolen
# z.B. "3 + 5 = 8" mit Unicode: ³ + ⁵ = ⁸ (hochgestellt)
# Oder LaTeX-ähnlich (falls QLabel HTML unterstützt)
''')

    doc.add_page_break()

    # === ABSCHNITT 6: TESTING ===
    add_heading(doc, '6. Testing-Strategie', 1)

    add_heading(doc, '6.1 Test-Dateien', 2)
    add_bullet_list(doc, [
        'tests/test_arithmetic_basic.py (~400 Zeilen)',
        'tests/test_arithmetic_concepts.py (~300 Zeilen)',
        'tests/test_arithmetic_advanced.py (~300 Zeilen)',
        'tests/test_number_language.py (~250 Zeilen)'
    ])

    add_heading(doc, '6.2 Test-Kategorien', 2)

    add_paragraph(doc, 'UNIT TESTS:', bold=True)
    add_bullet_list(doc, [
        'Einzelne Operationen (+, -, *, /)',
        'Vergleiche (<, >, =, etc.)',
        'Eigenschaften (gerade, prim, etc.)',
        'Zahl-Wort-Konvertierung',
        'Brüche und Dezimalzahlen'
    ])

    add_paragraph(doc, 'INTEGRATION TESTS:', bold=True)
    add_bullet_list(doc, [
        'End-to-End: UI Input → Berechnung → Output',
        'Neo4j Persistierung',
        'Proof Tree Generation',
        'Multi-Step Reasoning (z.B. "3+5 ist größer als 7?")'
    ])

    add_paragraph(doc, 'PROPERTY-BASED TESTS:', bold=True)
    add_bullet_list(doc, [
        'Hypothesis: Roundtrip parse/format',
        'Kommutativität: a+b = b+a',
        'Assoziativität: (a+b)+c = a+(b+c)',
        'Distributivität: a*(b+c) = a*b + a*c'
    ])

    add_heading(doc, '6.3 Test-Abdeckung Ziele', 2)
    add_bullet_list(doc, [
        'Line Coverage: >90%',
        'Branch Coverage: >85%',
        'Alle Edge Cases abgedeckt (Division durch 0, etc.)',
        'Performance: <100ms für einfache Operationen'
    ])

    doc.add_page_break()

    # === ABSCHNITT 7: VALIDIERUNG ===
    add_heading(doc, '7. Validierung & Qualitätssicherung', 1)

    add_heading(doc, '7.1 Funktionale Validierung', 2)
    add_numbered_list(doc, [
        'Alle Grundrechenarten funktionieren korrekt',
        'Zahl-Wort-Konvertierung bidirektional korrekt (0-999)',
        'Vergleiche produzieren korrekte Ergebnisse',
        'Eigenschaften (gerade, prim) werden korrekt erkannt',
        'Brüche werden automatisch gekürzt',
        'Proof Trees sind vollständig und nachvollziehbar',
        'Neo4j Integration funktioniert (Relationen gespeichert)',
        'UI zeigt Ergebnisse korrekt an'
    ])

    add_heading(doc, '7.2 Code-Qualität', 2)
    add_bullet_list(doc, [
        '✓ black + isort angewendet',
        '✓ flake8 bestanden (max-line-length=100)',
        '✓ mypy bestanden (type hints)',
        '✓ Alle Docstrings vorhanden',
        '✓ Keine pylint Warnungen (Score >9.0)'
    ])

    add_heading(doc, '7.3 Performance', 2)
    add_bullet_list(doc, [
        'Einfache Arithmetik: <50ms',
        'Komplexe Berechnungen: <200ms',
        'Zahl-Wort-Konvertierung: <10ms',
        'Primzahl-Check (n<10000): <100ms',
        'Neo4j Queries: <50ms'
    ])

    add_heading(doc, '7.4 Integration', 2)
    add_bullet_list(doc, [
        '✓ Keine Regression in bestehenden Tests',
        '✓ Alle test_*.py Dateien in settings_ui.py registriert',
        '✓ CLAUDE.md aktualisiert',
        '✓ Dokumentation vollständig (FEATURES_MATHEMATICS.md)',
        '✓ Commit-Messages folgen Konvention'
    ])

    doc.add_page_break()

    # === ABSCHNITT 8: CHECKLISTEN ===
    add_heading(doc, '8. Checklisten', 1)

    add_heading(doc, '8.1 Pre-Implementation Checklist', 2)
    add_code_block(doc, '''
□ Neo4j läuft und ist erreichbar
□ Alle Dependencies installiert (pytest, black, mypy, etc.)
□ Git Status sauber (keine uncommitted changes)
□ Bestehende Tests laufen alle durch (baseline)
□ Backup der Datenbank erstellt
□ Branch erstellt: feature/mathematics-module
''')

    add_heading(doc, '8.2 Phase 1 Checklist', 2)
    add_code_block(doc, '''
□ component_52_arithmetic_reasoning.py erstellt
□ ArithmeticEngine, BaseOperation, OperationRegistry implementiert
□ Addition, Subtraction, Multiplication, Division implementiert
□ component_53_number_language.py erstellt
□ NumberParser für 0-999 implementiert
□ NumberFormatter für 0-999 implementiert
□ Roundtrip-Tests bestanden
□ MeaningPoint.ARITHMETIC_QUESTION hinzugefügt
□ Intent Detection erweitert
□ GoalType.PERFORM_CALCULATION hinzugefügt
□ ArithmeticStrategy implementiert
□ ReasoningOrchestrator erweitert
□ tests/test_arithmetic_basic.py erstellt (20+ Tests)
□ Alle Tests bestehen
□ End-to-End Test erfolgreich: "Was ist drei plus fünf?"
□ Code formatiert (black + isort)
□ Commit: "feat: Add basic arithmetic reasoning"
''')

    add_heading(doc, '8.3 Phase 2 Checklist', 2)
    add_code_block(doc, '''
□ ComparisonEngine implementiert
□ Alle 5 Vergleichsoperatoren funktionieren
□ Transitive Inferenz implementiert
□ PropertyChecker implementiert (gerade, ungerade, prim)
□ find_divisors() implementiert
□ Eigenschaften in Neo4j gespeichert (HAS_PROPERTY)
□ Arithmetische Konzepte modelliert (Summe, Produkt, etc.)
□ tests/test_arithmetic_concepts.py erstellt
□ Alle Tests bestehen
□ Commit: "feat: Add arithmetic properties and concepts"
''')

    add_heading(doc, '8.4 Phase 3 Checklist', 2)
    add_code_block(doc, '''
□ RationalArithmetic implementiert
□ Brüche werden automatisch gekürzt
□ to_mixed_number() funktioniert
□ DecimalArithmetic implementiert
□ Rundung konfigurierbar
□ UI: Math Proof Tree Tab erstellt
□ Proof Trees werden korrekt angezeigt
□ Mathematische Symbole korrekt gerendert
□ tests/test_arithmetic_advanced.py erstellt
□ Alle Tests bestehen
□ Performance-Tests bestanden (<100ms)
□ Commit: "feat: Add rational/decimal arithmetic and UI"
''')

    add_heading(doc, '8.5 Documentation Checklist', 2)
    add_code_block(doc, '''
□ docs/FEATURES_MATHEMATICS.md erstellt
□ CLAUDE.md aktualisiert (Architektur-Sektion)
□ DEVELOPER_GUIDE.md erweitert (Math-Module)
□ Alle Komponenten dokumentiert (Docstrings)
□ Beispiele in Dokumentation eingefügt
□ README.md erweitert (neue Commands)
''')

    add_heading(doc, '8.6 Final Checklist', 2)
    add_code_block(doc, '''
□ Alle 38+ Tests bestehen (pytest tests/ -v)
□ Code Coverage >90% (pytest --cov)
□ black + isort angewendet
□ flake8 bestanden
□ mypy bestanden
□ settings_ui.py: Test-Dateien registriert
□ Neo4j Datenbank bereinigt (Test-Daten entfernt)
□ Performance-Tests bestanden
□ End-to-End Tests bestanden
□ Dokumentation vollständig
□ CLAUDE.md aktualisiert
□ Git: Feature-Branch in main gemerged
□ Final Commit: "feat: Complete mathematics module"
''')

    doc.add_page_break()

    # === ANHANG ===
    add_heading(doc, 'Anhang A: Datei-Übersicht', 1)

    add_heading(doc, 'NEUE DATEIEN:', 2)
    add_code_block(doc, '''
component_52_arithmetic_reasoning.py       ~1800 Zeilen
component_53_number_language.py            ~800 Zeilen
tests/test_arithmetic_basic.py             ~400 Zeilen
tests/test_arithmetic_concepts.py          ~300 Zeilen
tests/test_arithmetic_advanced.py          ~300 Zeilen
tests/test_number_language.py              ~250 Zeilen
docs/FEATURES_MATHEMATICS.md               ~200 Zeilen
                                           ──────────────
SUMME:                                     ~4050 Zeilen
''')

    add_heading(doc, 'MODIFIZIERTE DATEIEN:', 2)
    add_code_block(doc, '''
component_5_linguistik_strukturen.py       +5 Zeilen (MeaningPoint)
component_7_meaning_extractor.py           +80 Zeilen (Patterns)
component_4_goal_planner.py                +30 Zeilen (GoalType)
kai_sub_goal_executor.py                   +150 Zeilen (Strategy)
kai_reasoning_orchestrator.py              +50 Zeilen (Integration)
setup_initial_knowledge.py                 +20 Zeilen (Zahlen 1-10)
main_ui_graphical.py                       +30 Zeilen (UI)
settings_ui.py                             +5 Zeilen (Test Registry)
CLAUDE.md                                  +100 Zeilen (Doku)
                                           ─────────────
SUMME:                                     ~470 Zeilen
''')

    doc.add_page_break()

    add_heading(doc, 'Anhang B: Beispiel-Interaktionen', 1)

    add_paragraph(doc, 'BEISPIEL 1: Einfache Addition', bold=True)
    add_code_block(doc, '''
User: "Was ist drei plus fünf?"

KAI:
1. Meaning Extraction → ARITHMETIC_QUESTION (conf=0.95)
2. Number Parsing → [3, 5]
3. ArithmeticEngine.calculate("+", 3, 5) → 8
4. Number Formatting → "acht"
5. Response: "Die Summe von drei und fünf ist acht."

Proof Tree:
┌─ GIVEN: Operanden 3 und 5
│  └─ RULE_APPLICATION: Wende Addition an: 3 + 5
│     └─ CONCLUSION: Ergebnis: 8
''')

    add_paragraph(doc, 'BEISPIEL 2: Komplexe Frage mit Eigenschaften', bold=True)
    add_code_block(doc, '''
User: "Ist die Summe von vier und sechs eine gerade Zahl?"

KAI:
1. Erkenne: ARITHMETIC_QUESTION + PROPERTY_QUESTION
2. Berechne: 4 + 6 = 10
3. Prüfe Eigenschaft: is_even(10) → True
4. Response: "Ja, die Summe von vier und sechs ist zehn, und zehn ist eine gerade Zahl."

Proof Tree:
┌─ GIVEN: Operanden 4 und 6
│  └─ RULE_APPLICATION: Addition → 10
│     └─ RULE_APPLICATION: Prüfe 10 % 2 = 0
│        └─ CONCLUSION: 10 ist gerade
''')

    add_paragraph(doc, 'BEISPIEL 3: Bruchrechnung', bold=True)
    add_code_block(doc, '''
User: "Was ist sieben geteilt durch drei?"

KAI:
1. Number Parsing → [7, 3]
2. Division → Fraction(7, 3) = 2⅓
3. Response: "Sieben geteilt durch drei ist zwei Drittel (oder 2⅓)."
''')

    doc.add_page_break()

    add_heading(doc, 'Anhang C: Troubleshooting', 1)

    add_heading(doc, 'Problem: Zahl-Wort-Konvertierung fehlerhaft', 2)
    add_paragraph(doc, 'SYMPTOM: "einundzwanzig" → None oder falsche Zahl')
    add_paragraph(doc, 'LÖSUNG:')
    add_numbered_list(doc, [
        'Prüfe BASIC_NUMBERS und TENS Dictionaries',
        'Debug: print(parser.parse("einundzwanzig")) mit Zwischenschritten',
        'Teste Split-Logik: "einundzwanzig".split("und") → ["ein", "zwanzig"]',
        'Validiere Roundtrip: format(parse(word)) == word'
    ])

    add_heading(doc, 'Problem: Proof Tree nicht angezeigt', 2)
    add_paragraph(doc, 'SYMPTOM: UI zeigt keine Proof Trees an')
    add_paragraph(doc, 'LÖSUNG:')
    add_numbered_list(doc, [
        'Prüfe Signal-Verbindung in main_ui_graphical.py',
        'Validiere, dass ArithmeticResult.proof_tree nicht None ist',
        'Debug: print(proof_tree.to_dict()) vor UI-Update',
        'Prüfe ProofTreeWidget.display_proof() Implementation'
    ])

    add_heading(doc, 'Problem: Tests schlagen fehl', 2)
    add_paragraph(doc, 'SYMPTOM: pytest meldet Fehler')
    add_paragraph(doc, 'LÖSUNG:')
    add_numbered_list(doc, [
        'Isoliere fehlschlagende Tests: pytest tests/test_arithmetic_basic.py::test_name -v',
        'Prüfe Neo4j-Verbindung (läuft Datenbank?)',
        'Bereinige Test-Datenbank vor/nach Tests (fixtures)',
        'Prüfe Import-Errors: python -c "from component_52... import ..."'
    ])

    doc.add_page_break()

    # Schlusswort
    add_heading(doc, 'Schlusswort', 1)
    add_paragraph(doc, 'Diese Anleitung bietet eine detaillierte Roadmap für die Implementierung des '
                      'KAI Mathematik-Moduls. Folge den Schritten sequenziell und validiere nach jedem '
                      'Schritt die Funktionalität.')

    add_paragraph(doc, 'ERFOLGSFAKTOREN:', bold=True)
    add_bullet_list(doc, [
        'Inkrementelle Entwicklung (Schritt für Schritt)',
        'Tests nach jedem Schritt (sofortige Validierung)',
        'Klare Commit-Messages (nachvollziehbare Historie)',
        'Dokumentation parallel zur Implementierung',
        'Regelmäßiges Feedback (End-to-End Tests)'
    ])

    add_paragraph(doc, 'ZEITPLAN:', bold=True)
    add_bullet_list(doc, [
        'Phase 1: 3-4 Tage (Grundrechenarten + Integration)',
        'Phase 2: 2-3 Tage (Eigenschaften + Konzepte)',
        'Phase 3: 3-4 Tage (Brüche, Dezimalzahlen, UI)',
        'GESAMT: 8-11 Tage (inkl. Testing + Dokumentation)'
    ])

    add_paragraph(doc, '')
    add_paragraph(doc, 'Viel Erfolg bei der Implementierung!', bold=True, italic=True)
    add_paragraph(doc, '')
    add_paragraph(doc, '---')
    add_paragraph(doc, 'Erstellt: 2025-11-11')
    add_paragraph(doc, 'Version: 1.0')
    add_paragraph(doc, 'Projekt: KAI (Konzeptueller AI Prototyp)')

    return doc

# === MAIN ===
if __name__ == "__main__":
    print("Erstelle Mathematik-Modul Implementierungsanleitung...")
    doc = create_implementation_guide()

    output_path = "KAI_Mathematik_Modul_Implementierung.docx"
    doc.save(output_path)

    print(f"[OK] Dokument erstellt: {output_path}")
    print(f"[OK] Seiten: ~{len(doc.sections)} Abschnitte")
    print(f"[OK] Bereit zur Implementierung!")
